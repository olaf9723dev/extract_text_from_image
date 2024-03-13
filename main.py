from pytesseract import pytesseract
from PIL import Image
import cv2
import os
import math
import time
from docx import Document
from docx.shared import Pt, Mm
from docx2pdf import convert

class Extractor:
    def __init__(self) -> None:
        self.titles = []

    def write_docx(self):
        texts = self.get_texts()
        document = Document()
        self.set_layout(document)
        print('Create docx file.')
        for index, text in enumerate(texts):
            chapter = math.floor((index+2)/2)
            title = document.add_heading(f'Chapter {chapter} {self.titles[chapter-1]}', 0)
            self.set_font_style(title, 'Times New Roman', 18)
            p = document.add_paragraph(text)
            self.set_font_style(p, 'Times New Roman', 13)
            document.add_page_break()
            document.save('temp.docx')

    def set_font_style(self, text, style, size):
        run = text.runs[0]
        font = run.font
        font.name = style
        font.size = Pt(size)

    def get_texts(self):
        pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        paths = self.get_image_path()
        texts = []
        for index, path in enumerate(paths):
            print('Read Image ', index+1)
            image = cv2.imread(path)
            image= cv2.cvtColor(image,cv2.COLOR_BGR2RGB)
            # image = Image.open(path)
            text = pytesseract.image_to_string(image, timeout=5)
            if str(text.split('\n')[0]).isupper():
                self.titles.append(text.split('\n')[0])
            text = str(text).replace('\n', ' ')
            texts.append(text)
        return texts
    
    def get_image_path(self):
        print('Reading Image paths...')
        paths=[]
        for x in os.listdir('images'):
            if x.endswith(".jpg"):
               paths.append('images/'+ x)
        paths.sort()
        return paths
    
    def set_layout(self, document):
        section = document.sections[0]
        section.page_height = Mm(257)
        section.page_width = Mm(182)
        section.left_margin = Mm(30)
        section.right_margin = Mm(30)
        section.top_margin = Mm(30)
        section.bottom_margin = Mm(10)

    def convert_docx2pdf(self, input, output):
        print('Convert docx to Pdf file.')
        convert(input, output)

def main():
    start_time = time.time()
    extractor = Extractor()
    extractor.write_docx()
    extractor.convert_docx2pdf('temp.docx', 'temp.pdf')
    end_time = time.time()
    print('elapsed time : ', end_time-start_time, 's')
if __name__ == "__main__":
    main()